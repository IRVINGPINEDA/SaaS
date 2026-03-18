from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import csv
import re
import unicodedata


@dataclass
class StudentImportRow:
    matricula: str
    full_name: str
    category: str | None = None
    group_name: str | None = None
    password: str | None = None


def _norm(s: str) -> str:
    s = s.strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s)
    return s


def _guess_is_header(values: list[str]) -> bool:
    joined = " ".join(_norm(v) for v in values if v)
    return (
        "matricula" in joined
        or "nombre" in joined
        or "full name" in joined
        or "correo" in joined
        or "email" in joined
        or "grado" in joined
        or "carrera" in joined
    )


def _map_headers(headers: list[str]) -> dict[str, int]:
    h = [_norm(x) for x in headers]

    def find(*keys: str) -> int | None:
        for i, v in enumerate(h):
            for k in keys:
                if v == _norm(k) or _norm(k) in v:
                    return i
        return None

    idx_m = find("matricula", "matrícula", "id", "codigo", "código", "no control", "no. control")
    idx_n = find("nombre", "nombre completo", "alumno", "full_name", "full name", "fullname")
    # Categoria: a veces es "carrera" o "programa"
    idx_c = find("categoria", "categoría", "category", "carrera", "programa", "especialidad")
    # Grupo: a veces es "grado" o "semestre"
    idx_g = find("grupo", "group", "grado", "semestre", "seccion", "sección")
    idx_p = find("password", "contraseña", "contrasena", "clave")

    mapping: dict[str, int] = {}
    if idx_m is not None:
        mapping["matricula"] = idx_m
    if idx_n is not None:
        mapping["full_name"] = idx_n
    if idx_c is not None:
        mapping["category"] = idx_c
    if idx_g is not None:
        mapping["group_name"] = idx_g
    if idx_p is not None:
        mapping["password"] = idx_p
    return mapping


def _rows_from_delimited_text(text: str) -> list[list[str]]:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    rows: list[list[str]] = []
    for ln in lines:
        # Separadores comunes: ; , | \t
        parts = re.split(r"\s*[;,|\t]\s*", ln)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 2:
            rows.append(parts)
            continue

        # Tablas en PDF suelen venir por espacios multiples entre columnas
        parts2 = re.split(r"\s{2,}", ln)
        parts2 = [p.strip() for p in parts2 if p.strip()]
        if len(parts2) >= 2:
            rows.append(parts2)
            continue

        # Heuristica: lineas que comienzan con matricula y contienen email/telefono
        parsed = _parse_line_row(ln)
        if parsed:
            rows.append(parsed)
    return rows


_MAT_RE = re.compile(r"^([A-Za-z0-9]{1,6}\d{0,4}-\d{3,6}|[A-Za-z0-9-]{6,})\s+(.+)$")
_EMAIL_RE = re.compile(r"[\w.\-+]+@[\w.\-]+\.\w+")
_PHONE_RE = re.compile(r"(\+?\d[\d\-\s]{7,}\d)")


def _parse_line_row(line: str) -> list[str] | None:
    m = _MAT_RE.match(line.strip())
    if not m:
        return None
    matricula = m.group(1).strip()
    rest = m.group(2).strip()

    # extraer email y telefono (si existen)
    email = None
    phone = None

    em = _EMAIL_RE.search(rest)
    if em:
        email = em.group(0)
        rest = (rest[: em.start()] + " " + rest[em.end() :]).strip()

    ph = _PHONE_RE.search(rest)
    if ph:
        phone = ph.group(0).strip()
        rest = (rest[: ph.start()] + " " + rest[ph.end() :]).strip()

    tokens = [t for t in rest.split() if t]
    if len(tokens) < 1:
        return None

    # buscar edad (primer entero razonable)
    age_idx = None
    for i, t in enumerate(tokens):
        if t.isdigit():
            try:
                v = int(t)
            except Exception:
                continue
            if 10 <= v <= 99:
                age_idx = i
                break

    full_name = ""
    group_name = None
    category = None

    if age_idx is None:
        full_name = " ".join(tokens).strip()
    else:
        full_name = " ".join(tokens[:age_idx]).strip()
        # siguiente token suele ser grado/grupo (ej. 5°)
        if age_idx + 1 < len(tokens):
            group_name = tokens[age_idx + 1].strip()
        # el resto suele ser carrera/categoria
        tail_start = age_idx + 2
        if tail_start < len(tokens):
            category = " ".join(tokens[tail_start:]).strip() or None

    if not full_name:
        return None

    # devolvemos columnas basicas: matricula, nombre, categoria, grupo, email, telefono
    out = [matricula, full_name]
    if category:
        out.append(category)
    if group_name:
        # garantizamos que haya al menos 3 columnas para que _parse_rows pueda tomar category como col 2
        if len(out) < 3:
            out.append("")
        out.append(group_name)
    if email:
        out.append(email)
    if phone:
        out.append(phone)
    return out


def parse_csv(data: bytes) -> list[StudentImportRow]:
    try:
        text = data.decode("utf-8-sig")
    except Exception:
        text = data.decode("latin-1", errors="ignore")

    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except Exception:
        dialect = csv.excel

    reader = csv.reader(text.splitlines(), dialect)
    raw = [row for row in reader if any((c or "").strip() for c in row)]
    return _parse_rows(raw)


def parse_xlsx(data: bytes) -> list[StudentImportRow]:
    # openpyxl es dependencia opcional del import
    from openpyxl import load_workbook

    wb = load_workbook(filename=BytesIO(data), data_only=True)
    ws = wb.active

    raw: list[list[str]] = []
    for row in ws.iter_rows(values_only=True):
        values = ["" if v is None else str(v).strip() for v in row]
        if not any(values):
            continue
        raw.append(values)

    return _parse_rows(raw)


def parse_docx(data: bytes) -> list[StudentImportRow]:
    # python-docx
    from docx import Document

    doc = Document(BytesIO(data))

    raw: list[list[str]] = []

    # Tablas primero (más estructurado)
    for table in doc.tables:
        for row in table.rows:
            vals = [(_cell.text or "").strip() for _cell in row.cells]
            if any(vals):
                raw.append(vals)

    if raw:
        return _parse_rows(raw)

    # Fallback: párrafos con separadores
    text = "\n".join((p.text or "") for p in doc.paragraphs)
    return _parse_rows(_rows_from_delimited_text(text))


def parse_pdf(data: bytes) -> list[StudentImportRow]:
    # pypdf (extracción de texto)
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    text = ""
    for page in reader.pages:
        t = page.extract_text() or ""
        text += t + "\n"
    raw = _rows_from_delimited_text(text)
    return _parse_rows(raw)


def _parse_rows(raw: list[list[str]]) -> list[StudentImportRow]:
    if not raw:
        return []

    # Detect header
    first = [str(x or "").strip() for x in raw[0]]
    has_header = _guess_is_header(first)

    mapping: dict[str, int] = {}
    start = 0
    if has_header:
        mapping = _map_headers(first)
        start = 1

    rows: list[StudentImportRow] = []
    for r in raw[start:]:
        vals = [str(x or "").strip() for x in r]
        vals = [v for v in vals if v != ""]
        if len(vals) < 2:
            continue

        def get(key: str, fallback_index: int | None) -> str | None:
            idx = mapping.get(key)
            if idx is not None and idx < len(r):
                v = str(r[idx] or "").strip()
                return v or None
            if idx is None and fallback_index is not None and fallback_index < len(vals):
                v = vals[fallback_index].strip()
                return v or None
            return None

        # Fallback por posiciones si no hay headers:
        # - Caso simple: [matricula, nombre, carrera/categoria, grado/grupo]
        # - Caso comun en listas: [matricula, nombre, edad, grado, carrera, ...]
        if not mapping and len(vals) >= 5 and vals[2].isdigit():
            matricula = vals[0]
            full_name = vals[1]
            group_name = vals[3] if len(vals) > 3 else None
            category = vals[4] if len(vals) > 4 else None
            password = vals[5] if len(vals) > 5 else None
        else:
            matricula = get("matricula", 0) or ""
            full_name = get("full_name", 1) or ""
            category = get("category", 2)
            group_name = get("group_name", 3)
            password = get("password", 4)

        matricula = matricula.strip()
        full_name = full_name.strip()
        if not matricula or not full_name:
            continue

        rows.append(
            StudentImportRow(
                matricula=matricula,
                full_name=full_name,
                category=(category.strip() if category else None),
                group_name=(group_name.strip() if group_name else None),
                password=(password.strip() if password else None),
            )
        )

    return rows
